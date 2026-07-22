# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x1F, 0x2A, 0x37)
ACCENT = RGBColor(0x2E, 0x6B, 0xE0)
GRAY = RGBColor(0x55, 0x5F, 0x6B)

def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def style_header_row(row):
    for cell in row.cells:
        set_cell_shading(cell, "2E6BE0")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)

def add_heading(doc, text, size=16, color=DARK, space_before=18, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.color.rgb = color
    return p

def add_body(doc, text, size=10.5, color=DARK, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    return p

def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = DARK
    return p

def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    style_header_row(hdr)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.color.rgb = DARK
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


doc = Document()

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# Margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ---------- Portada / título ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("GymAccess")
r.font.size = Pt(28)
r.font.bold = True
r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Acuerdo de Servicio y Cotización")
r.font.size = Pt(14)
r.font.color.rgb = GRAY

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Sistema de gestión y control de acceso para gimnasios")
r.font.size = Pt(10.5)
r.italic = True
r.font.color.rgb = GRAY

doc.add_paragraph()

# ---------- Datos del acuerdo ----------
add_heading(doc, "Datos del acuerdo", size=13)
make_table(
    doc,
    headers=["Campo", "Detalle"],
    rows=[
        ["Cliente / Gimnasio", "[NOMBRE DEL GIMNASIO]"],
        ["Representante del cliente", "[NOMBRE Y CARGO]"],
        ["Proveedor", "Adrian Valencia — GymAccess"],
        ["Contacto proveedor", "3322332046 · angel10000sv@gmail.com"],
        ["Fecha", "[DD / MM / 2026]"],
        ["Plan contratado", "[Plan Arranque / Plan Crecimiento]"],
    ],
    col_widths=[5.5, 10.5],
)

# ---------- Objeto del acuerdo ----------
add_heading(doc, "1. Objeto del acuerdo")
add_body(
    doc,
    "El proveedor se compromete a instalar, configurar y dar soporte al sistema GymAccess "
    "en las instalaciones del cliente, conforme al plan y equipo detallados en las tablas de "
    "este documento. El cliente se compromete a realizar los pagos en los términos aquí descritos."
)

# ---------- Cláusula de mes de prueba ----------
add_heading(doc, "2. Periodo de prueba — primer mes sin costo")
add_body(
    doc,
    "El primer mes de suscripción NO se cobra. Este periodo se considera de puesta a punto y "
    "validación en un entorno real de operación del gimnasio."
)
add_bullet(doc, "Cualquier ajuste, corrección o falla detectada durante este mes se resuelve sin costo adicional para el cliente.")
add_bullet(doc, "El objetivo es entregar el sistema completamente pulido y estable antes de iniciar el cobro de la mensualidad.")
add_bullet(doc, "El cobro de la suscripción mensual inicia a partir del segundo mes de uso, conforme al plan contratado.")
add_bullet(doc, "El pago único de equipo e instalación (si aplica) se mantiene conforme a lo indicado en la tabla de inversión inicial.")

# ---------- Plan y suscripción ----------
add_heading(doc, "3. Planes de suscripción")
make_table(
    doc,
    headers=["Concepto", "Plan Arranque", "Plan Crecimiento"],
    rows=[
        ["Socios activos", "Hasta 200", "Hasta 600 en total"],
        ["Sucursales", "1", "Hasta 3"],
        ["Empleados (recepción, admin, coaches)", "Hasta 20", "Hasta 20 por sucursal"],
        ["Control de acceso por QR", "Incluido", "Incluido"],
        ["Pagos, cobranza y reportes", "Incluido", "Incluido"],
        ["Dashboard por sucursal", "—", "Incluido"],
        ["Soporte", "Correo / WhatsApp", "Correo / WhatsApp prioritario"],
        ["Precio mensual (a partir del mes 2)", "$599 MXN / mes", "$1,199 MXN / mes"],
    ],
    col_widths=[6.5, 5, 5],
)
add_body(doc, "Plan seleccionado por el cliente: [Plan Arranque / Plan Crecimiento]", bold=True)

add_heading(doc, "4. Expansiones opcionales", size=12)
make_table(
    doc,
    headers=["Expansión", "Precio"],
    rows=[
        ["Paquete de 50 socios extra", "+$100 MXN / mes"],
        ["Sucursal extra (incluye 200 socios adicionales)", "+$300 MXN / mes"],
    ],
    col_widths=[10.5, 5.5],
)
add_body(doc, "Sin plazos forzosos: el cliente puede activar o desactivar expansiones cuando lo requiera.", size=9.5, color=GRAY)

# ---------- Equipo ----------
add_heading(doc, "5. Equipo necesario (inversión única, por sucursal)")
make_table(
    doc,
    headers=["Equipo", "Cantidad", "Precio aprox. (MXN)"],
    rows=[
        ["Mini PC o PC seminueva", "1", "$2,000 – $2,300"],
        ["Teclado y mouse", "1", "$424"],
        ["Monitor", "1", "$900 – $1,600"],
        ["Lector de código QR USB", "1", "$399 – $800"],
        ["Total equipo", "", "$3,700 – $5,200"],
    ],
    col_widths=[9, 3, 4],
)
add_body(doc, "Precios de referencia (Mercado Libre / Acteck / Facebook Marketplace). El equipo es propiedad del cliente, pago único. Si el cliente ya cuenta con computadora, solo se requiere el lector QR (~$500).", size=9.5, color=GRAY)

# ---------- Inversión inicial ----------
add_heading(doc, "6. Inversión inicial (pago único)")
make_table(
    doc,
    headers=["Concepto", "Precio (MXN)"],
    rows=[
        ["Instalación y configuración inicial", "$2,000"],
        ["Capacitación del personal", "Incluida en la instalación"],
        ["Migración de socios actuales", "Incluida en la instalación"],
    ],
    col_widths=[10.5, 5.5],
)

# ---------- Resumen de costos ----------
add_heading(doc, "7. Resumen de costos")
make_table(
    doc,
    headers=["Concepto", "Momento del pago", "Monto (MXN)"],
    rows=[
        ["Equipo (inversión única)", "Antes de instalación", "$3,700 – $5,200"],
        ["Instalación y configuración", "Antes de instalación", "$2,000"],
        ["Mes 1 de suscripción", "Mes 1", "$0 (cortesía / periodo de prueba)"],
        ["Mes 2 en adelante", "Mensual", "$599 (Arranque) / $1,199 (Crecimiento)"],
    ],
    col_widths=[6, 5, 5],
)

# ---------- Requisitos ----------
add_heading(doc, "8. Requisitos del gimnasio")
add_bullet(doc, "Internet estable (10 Mbps es suficiente).")
add_bullet(doc, "Un navegador web (Chrome, Edge, etc.).")
add_bullet(doc, "El proveedor se encarga del servidor, respaldos y seguridad del sistema.")

# ---------- Siguientes pasos ----------
add_heading(doc, "9. Siguientes pasos")
add_bullet(doc, "Confirmar el plan y adquirir el equipo.")
add_bullet(doc, "Instalación y configuración del sistema (1 día).")
add_bullet(doc, "Capacitación del personal.")
add_bullet(doc, "Migración de socios actuales.")
add_bullet(doc, "Inicio del mes de prueba sin costo.")

# ---------- Firmas ----------
add_heading(doc, "10. Conformidad")
add_body(doc, "Ambas partes firman de conformidad con los términos aquí descritos.", space_after=24)

sig_table = doc.add_table(rows=2, cols=2)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sig_table.cell(0, 0).text = "_____________________________"
sig_table.cell(0, 1).text = "_____________________________"
sig_table.cell(1, 0).text = "Por el Cliente — [NOMBRE]"
sig_table.cell(1, 1).text = "Por el Proveedor — Adrian Valencia"
for row in sig_table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(10)

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("GymAccess · Adrian Valencia · 3322332046 · angel10000sv@gmail.com")
r.font.size = Pt(9)
r.font.color.rgb = GRAY

doc.save(r"C:\Users\Angel.sanchez\Desktop\CLAUDE\gymAccess\cotizacion\GymAccess_Acuerdo_Cotizacion.docx")
print("done")
