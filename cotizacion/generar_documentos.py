# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r"C:\Users\Angel.sanchez\Desktop\CLAUDE\gymAccess\cotizacion"

CLIENTE = "Round3Boxing"
REP_CLIENTE = "Diego Sandoval"
PROVEEDOR = "Adrian Valencia"
MARCA = "GymAccess"
CONTACTO = "3322332046 · angel10000sv@gmail.com"
FECHA = "25/07/2026"
ESTADO = "Jalisco, México"  # ajustar si no aplica

def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def style_doc(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(MARCA)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x57)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(title)
    run2.bold = True
    run2.font.size = Pt(14)

    if subtitle:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(subtitle)
        run3.italic = True
        run3.font.size = Pt(10)
        run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x57)
    return h

def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

def add_datos_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in rows:
        row = table.add_row()
        row.cells[0].text = k
        row.cells[1].text = v
        row.cells[0].paragraphs[0].runs[0].bold = True
    return table

def add_signature_block(doc, cliente_label=f"Por el Cliente — {REP_CLIENTE}", proveedor_label=f"Por el Proveedor — {PROVEEDOR}"):
    doc.add_paragraph()
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=2)
    table.autofit = True
    table.cell(0, 0).text = "_____________________________"
    table.cell(0, 1).text = "_____________________________"
    table.cell(1, 0).text = cliente_label
    table.cell(1, 1).text = proveedor_label
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_footer_line(doc):
    doc.add_paragraph()
    p = add_para(doc, f"{MARCA} · {PROVEEDOR} · {CONTACTO}", italic=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# =========================================================
# 1. AVISO DE PRIVACIDAD
# =========================================================
def gen_aviso_privacidad():
    doc = Document()
    style_doc(doc)
    add_title(doc, "Aviso de Privacidad", "Sistema de gestión y control de acceso para gimnasios")

    add_para(doc,
        f"En cumplimiento con la Ley General de Protección de Datos Personales en Posesión de Sujetos "
        f"Obligados y la Ley Federal de Protección de Datos Personales en Posesión de los Particulares "
        f"(México), {PROVEEDOR}, operando bajo la marca {MARCA} (en adelante \"el Proveedor\"), pone a "
        f"disposición de {CLIENTE} (en adelante \"el Responsable\") y de los titulares de datos personales "
        f"(socios, empleados y personal del gimnasio) el presente Aviso de Privacidad.")

    add_heading(doc, "1. Identidad y domicilio del responsable")
    add_para(doc,
        f"{CLIENTE}, a través de su representante {REP_CLIENTE}, es el Responsable del tratamiento de los "
        f"datos personales de sus socios y empleados. {MARCA} ({PROVEEDOR}) actúa como Encargado del "
        f"tratamiento, proporcionando la plataforma tecnológica que permite dicho tratamiento por cuenta "
        f"del Responsable, conforme al Contrato de Prestación de Servicios y al Acuerdo de Encargado de "
        f"Datos vigente entre las partes.")

    add_heading(doc, "2. Datos personales recabados")
    add_para(doc, "Para la prestación del servicio se recaban, entre otros, los siguientes datos:")
    add_bullets(doc, [
        "Datos de identificación: nombre completo, teléfono, correo electrónico, fotografía (opcional).",
        "Datos de membresía: plan contratado, fecha de inscripción, historial de asistencia.",
        "Datos de pago: monto, fecha y método de pago (no se almacenan datos completos de tarjeta cuando el cobro se procesa por un tercero certificado).",
        "Dato biométrico sensible: huella dactilar, utilizada exclusivamente para control de acceso mediante lector QR/biométrico.",
    ])
    add_para(doc,
        "El dato biométrico (huella dactilar) es considerado un dato personal sensible. Su recolección "
        "requiere el consentimiento expreso y por escrito del titular, mismo que el Responsable (el "
        "gimnasio) debe recabar previo al registro del socio o empleado en el sistema.")

    add_heading(doc, "3. Finalidades del tratamiento")
    add_para(doc, "Los datos se utilizan para las siguientes finalidades necesarias para el servicio:")
    add_bullets(doc, [
        "Control de acceso e identificación de socios en las instalaciones.",
        "Gestión de membresías, cobros y renovaciones.",
        "Generación de reportes administrativos para el gimnasio.",
        "Comunicación con el socio sobre su membresía (recordatorios, vencimientos).",
    ])
    add_para(doc, "No se utilizan los datos para finalidades secundarias (publicidad, venta a terceros, mercadotecnia) sin consentimiento adicional expreso del titular.")

    add_heading(doc, "4. Transferencia de datos")
    add_para(doc,
        "Los datos no se transfieren a terceros ajenos a la operación del servicio, salvo: (a) proveedores "
        "de infraestructura tecnológica (hosting/servidor) necesarios para operar el sistema, bajo obligación "
        "contractual de confidencialidad y seguridad; (b) autoridad competente cuando exista requerimiento legal.")

    add_heading(doc, "5. Derechos ARCO")
    add_para(doc,
        "El titular de los datos (socio o empleado) tiene derecho a Acceder, Rectificar, Cancelar u Oponerse "
        "(derechos ARCO) al tratamiento de sus datos personales, así como a revocar su consentimiento en "
        "cualquier momento. Para ejercer estos derechos, el titular debe dirigirse directamente al gimnasio "
        f"({CLIENTE}), quien es el Responsable del tratamiento, a través de su representante {REP_CLIENTE}.")

    add_heading(doc, "6. Medidas de seguridad")
    add_para(doc, "El Proveedor implementa medidas técnicas y administrativas para proteger los datos personales, incluyendo:")
    add_bullets(doc, [
        "Acceso restringido al servidor y base de datos, limitado al personal autorizado.",
        "Respaldos periódicos de la información.",
        "Cifrado de contraseñas y datos sensibles en almacenamiento cuando la plataforma lo soporte.",
        "Notificación al Responsable en caso de vulneración de seguridad que afecte datos personales, a la brevedad posible.",
    ])

    add_heading(doc, "7. Conservación de datos")
    add_para(doc,
        "Los datos personales se conservan mientras dure la relación de membresía del socio con el gimnasio "
        "y durante el tiempo adicional necesario para cumplir obligaciones legales o contractuales. En caso "
        "de terminación del contrato entre el Responsable y el Proveedor, los datos se entregan al "
        "Responsable y se eliminan de los sistemas del Proveedor conforme a lo pactado en el Contrato de "
        "Prestación de Servicios.")

    add_heading(doc, "8. Cambios al aviso de privacidad")
    add_para(doc,
        "Cualquier modificación a este aviso será comunicada a través de los medios de contacto registrados "
        "o mediante publicación visible en las instalaciones del gimnasio o en la plataforma.")

    add_heading(doc, "9. Contacto")
    add_datos_table(doc, [
        ("Responsable (gimnasio)", f"{CLIENTE} — {REP_CLIENTE}"),
        ("Encargado (proveedor tecnológico)", f"{MARCA} — {PROVEEDOR}"),
        ("Contacto encargado", CONTACTO),
        ("Fecha de emisión", FECHA),
    ])

    add_footer_line(doc)
    doc.save(f"{OUT_DIR}\\GymAccess_Aviso_de_Privacidad.docx")
    print("Aviso de Privacidad generado.")


# =========================================================
# 2. ACUERDO DE CONFIDENCIALIDAD (NDA)
# =========================================================
def gen_nda():
    doc = Document()
    style_doc(doc)
    add_title(doc, "Acuerdo de Confidencialidad", "(NDA — Non-Disclosure Agreement)")

    add_para(doc,
        f"El presente Acuerdo de Confidencialidad se celebra entre {CLIENTE}, representado por "
        f"{REP_CLIENTE} (en adelante \"el Cliente\"), y {PROVEEDOR}, operando bajo la marca {MARCA} "
        f"(en adelante \"el Proveedor\"), en conjunto referidos como \"las Partes\", con el objeto de "
        f"proteger la información confidencial que se compartan mutuamente con motivo de la prestación "
        f"del servicio de gestión y control de acceso para gimnasios.")

    add_heading(doc, "1. Definición de información confidencial")
    add_para(doc, "Se considera información confidencial, de forma enunciativa mas no limitativa:")
    add_bullets(doc, [
        "Del Cliente: base de datos de socios, información financiera y de cobranza, procesos internos de operación del gimnasio, listas de precios y promociones internas.",
        "Del Proveedor: código fuente, arquitectura del sistema, metodologías de desarrollo, estructura de precios y condiciones comerciales no publicadas.",
    ])

    add_heading(doc, "2. Obligaciones de las partes")
    add_bullets(doc, [
        "Utilizar la información confidencial únicamente para los fines relacionados con la prestación del servicio.",
        "No divulgar, publicar, ni compartir la información confidencial con terceros sin autorización previa y por escrito de la otra parte.",
        "Aplicar medidas razonables de seguridad para resguardar la información confidencial a la que tengan acceso.",
        "Limitar el acceso a la información confidencial únicamente al personal que la requiera para cumplir con el objeto del servicio.",
    ])

    add_heading(doc, "3. Excepciones")
    add_para(doc, "No se considera incumplimiento cuando la información:")
    add_bullets(doc, [
        "Sea de dominio público sin culpa de la parte receptora.",
        "Deba divulgarse por requerimiento de autoridad competente, notificando previamente a la otra parte cuando sea legalmente posible.",
        "Ya fuera conocida por la parte receptora antes de la firma del presente acuerdo.",
    ])

    add_heading(doc, "4. Vigencia")
    add_para(doc,
        "Este acuerdo entra en vigor a partir de la fecha de firma y permanece vigente durante toda la "
        "relación comercial entre las Partes, y se extiende por un periodo de 2 (dos) años posteriores a "
        "la terminación del Contrato de Prestación de Servicios, respecto de la información confidencial "
        "compartida durante la vigencia de dicha relación.")

    add_heading(doc, "5. Propiedad de la información")
    add_para(doc,
        "Ninguna disposición de este acuerdo se interpreta como transferencia de derechos de propiedad "
        "intelectual o de datos. Cada parte conserva la titularidad de su propia información confidencial.")

    add_heading(doc, "6. Jurisdicción")
    add_para(doc,
        f"Para la interpretación y cumplimiento del presente acuerdo, las Partes se someten a las leyes "
        f"aplicables en {ESTADO}, renunciando a cualquier otro fuero que pudiera corresponderles por razón "
        f"de su domicilio presente o futuro.")

    add_datos_table(doc, [
        ("Cliente", f"{CLIENTE} — {REP_CLIENTE}"),
        ("Proveedor", f"{MARCA} — {PROVEEDOR}"),
        ("Contacto proveedor", CONTACTO),
        ("Fecha", FECHA),
    ])

    add_heading(doc, "Conformidad")
    add_para(doc, "Ambas partes firman de conformidad con los términos aquí descritos.")
    add_signature_block(doc)
    add_footer_line(doc)
    doc.save(f"{OUT_DIR}\\GymAccess_Acuerdo_Confidencialidad.docx")
    print("Acuerdo de Confidencialidad generado.")


# =========================================================
# 3. CONTRATO DE PRESTACIÓN DE SERVICIOS
# =========================================================
def gen_contrato_servicios():
    doc = Document()
    style_doc(doc)
    add_title(doc, "Contrato de Prestación de Servicios", "Sistema de gestión y control de acceso para gimnasios")

    add_para(doc,
        f"Contrato que celebran, por una parte {CLIENTE}, representado por {REP_CLIENTE} (en adelante "
        f"\"el Cliente\"), y por otra parte {PROVEEDOR}, operando bajo la marca {MARCA} (en adelante "
        f"\"el Proveedor\"), al tenor de las siguientes cláusulas.")

    add_heading(doc, "1. Objeto del contrato")
    add_para(doc,
        "El Proveedor se compromete a instalar, configurar, dar mantenimiento y soporte al sistema "
        "GymAccess en las instalaciones del Cliente, conforme al plan de suscripción, equipo y condiciones "
        "económicas establecidas en la Cotización firmada entre las partes, misma que forma parte integral "
        "de este contrato como Anexo A.")

    add_heading(doc, "2. Vigencia")
    add_para(doc,
        "El presente contrato tiene vigencia indefinida a partir de la fecha de firma, sujeto a renovación "
        "mensual automática conforme al ciclo de facturación del plan contratado, salvo terminación "
        "conforme a la cláusula 6.")

    add_heading(doc, "3. Propiedad intelectual")
    add_para(doc,
        "El código fuente, arquitectura, y demás elementos de propiedad intelectual del sistema GymAccess "
        "son y permanecen propiedad exclusiva del Proveedor. El Cliente adquiere únicamente una licencia "
        "de uso del sistema durante la vigencia de este contrato, no transferible, para los fines de "
        "operación de su gimnasio.")
    add_para(doc,
        "El equipo físico (mini PC, monitor, lector QR, periféricos) es propiedad del Cliente una vez "
        "cubierto su pago conforme al Anexo A.")

    add_heading(doc, "4. Propiedad y tratamiento de datos")
    add_para(doc,
        "Los datos personales de socios y empleados capturados en el sistema (incluyendo datos "
        "biométricos como huella dactilar) son y permanecen propiedad del Cliente, quien funge como "
        "Responsable del tratamiento de dichos datos conforme al Aviso de Privacidad vigente.")
    add_para(doc,
        "El Proveedor funge como Encargado del tratamiento y se obliga a:")
    add_bullets(doc, [
        "Tratar los datos personales únicamente conforme a las instrucciones del Cliente y para los fines del servicio contratado.",
        "No utilizar los datos del Cliente para fines distintos a la prestación del servicio, ni compartirlos con terceros no autorizados.",
        "Implementar medidas de seguridad técnicas y administrativas razonables (respaldos, control de acceso al servidor, cifrado cuando aplique).",
        "Notificar al Cliente en un plazo no mayor a 72 horas ante cualquier incidente de seguridad que comprometa los datos personales.",
        "A la terminación del contrato, entregar al Cliente una copia completa de su base de datos y eliminar dicha información de los sistemas del Proveedor en un plazo de 30 días naturales, salvo obligación legal de conservación.",
    ])

    add_heading(doc, "5. Nivel de servicio y soporte (SLA)")
    add_para(doc, "El Proveedor ofrece soporte técnico conforme al plan contratado, con los siguientes tiempos de respuesta:")
    add_bullets(doc, [
        "Incidencias críticas (sistema caído, sin control de acceso): respuesta en un máximo de 4 horas hábiles.",
        "Incidencias de prioridad alta (falla parcial, error en cobros/reportes): respuesta en un máximo de 24 horas hábiles.",
        "Incidencias de prioridad media/baja (dudas, ajustes menores): respuesta en un máximo de 48 horas hábiles.",
        "Canal de soporte: WhatsApp/correo indicados en este contrato, en horario de 9:00 a 20:00 hrs, de lunes a sábado.",
    ])
    add_para(doc,
        "\"Respuesta\" implica confirmación de recepción del reporte y diagnóstico inicial, no necesariamente "
        "la resolución total del problema en dicho plazo, misma que dependerá de la complejidad de la falla.")

    add_heading(doc, "6. Terminación y cancelación")
    add_para(doc, "Cualquiera de las partes puede dar por terminado este contrato bajo las siguientes condiciones:")
    add_bullets(doc, [
        "Aviso previo: la parte que desee terminar el contrato debe notificar a la otra con al menos 30 días naturales de anticipación.",
        "El Cliente puede cancelar el servicio de suscripción mensual en cualquier momento posterior al mes de prueba, sin penalización, notificando conforme al plazo anterior.",
        "El equipo adquirido por el Cliente (pagado conforme al Anexo A) es de su propiedad y no es reembolsable una vez comprado e instalado, dado que constituye una inversión única e independiente de la suscripción mensual.",
        "En caso de cancelación durante el proceso de instalación (posterior al pago del anticipo y previo a la entrega del equipo), el Proveedor reembolsará el anticipo únicamente descontando los costos ya erogados en la adquisición del equipo (comprobables mediante factura o comprobante de compra).",
        "El Proveedor puede suspender o terminar el servicio por falta de pago de la suscripción mensual, previa notificación al Cliente con 5 días naturales para regularizar el pago antes de la suspensión.",
        "A la terminación del contrato por cualquier causa, aplica lo establecido en la cláusula 4 respecto a la entrega y eliminación de datos.",
    ])

    add_heading(doc, "7. Responsabilidad")
    add_para(doc,
        "El Proveedor no será responsable por fallas derivadas de causas ajenas a su control, incluyendo "
        "fallas de conectividad a internet, suministro eléctrico, o mal uso del equipo por parte del "
        "personal del Cliente. El Proveedor es responsable de la operación, mantenimiento, seguridad y "
        "respaldo del servidor y del sistema.")

    add_heading(doc, "8. Jurisdicción y ley aplicable")
    add_para(doc,
        f"Para la interpretación y cumplimiento del presente contrato, así como para todo aquello que no "
        f"esté expresamente previsto en el mismo, las Partes se someten a la legislación aplicable en "
        f"{ESTADO}, y a los tribunales competentes de dicha jurisdicción, renunciando expresamente a "
        f"cualquier otro fuero que pudiera corresponderles por razón de su domicilio presente o futuro.")

    add_heading(doc, "9. Documentos que forman parte del contrato")
    add_bullets(doc, [
        "Anexo A: Cotización de Servicio (GymAccess_Acuerdo_Cotizacion.docx).",
        "Anexo B: Aviso de Privacidad (GymAccess_Aviso_de_Privacidad.docx).",
        "Anexo C: Acuerdo de Confidencialidad (GymAccess_Acuerdo_Confidencialidad.docx).",
    ])

    add_datos_table(doc, [
        ("Cliente", f"{CLIENTE} — {REP_CLIENTE}"),
        ("Proveedor", f"{MARCA} — {PROVEEDOR}"),
        ("Contacto proveedor", CONTACTO),
        ("Fecha", FECHA),
    ])

    add_heading(doc, "Conformidad")
    add_para(doc, "Ambas partes firman de conformidad con los términos aquí descritos.")
    add_signature_block(doc)
    add_footer_line(doc)
    doc.save(f"{OUT_DIR}\\GymAccess_Contrato_Prestacion_Servicios.docx")
    print("Contrato de Prestacion de Servicios generado.")


# =========================================================
# 4. FORMATO DE INSCRIPCIÓN Y CONSENTIMIENTO DE DATOS (SOCIO)
# =========================================================
def gen_consentimiento_socio():
    doc = Document()
    style_doc(doc)
    add_title(doc, "Formato de Inscripción y Consentimiento de Datos", f"Uso exclusivo de {CLIENTE}")

    add_para(doc,
        "Este formato debe ser completado y firmado por cada socio al momento de su inscripción o "
        "renovación de membresía, para el registro de sus datos personales en el sistema de control "
        "de acceso del gimnasio.")

    add_heading(doc, "1. Datos del socio")
    add_datos_table(doc, [
        ("Nombre completo", "_______________________________________________"),
        ("Fecha de nacimiento", "___ /___ /______"),
        ("Teléfono de contacto", "_______________________________________________"),
        ("Correo electrónico", "_______________________________________________"),
        ("Plan / membresía", "_______________________________________________"),
        ("Fecha de inscripción", "___ /___ /______"),
    ])

    add_heading(doc, "2. Consentimiento de tratamiento de datos personales")
    add_para(doc,
        f"Declaro que he sido informado(a) del Aviso de Privacidad de {CLIENTE}, y autorizo el "
        f"tratamiento de mis datos personales (nombre, contacto, historial de membresía y pagos) "
        f"para fines de administración de mi membresía, control de acceso y comunicación relacionada "
        f"con el servicio del gimnasio.")
    add_para(doc, "Acepto:   [  ] Sí       [  ] No")

    add_heading(doc, "3. Consentimiento expreso — dato biométrico (huella dactilar)")
    add_para(doc,
        "Autorizo de manera expresa el uso de mi huella dactilar exclusivamente para fines de control "
        "de acceso a las instalaciones del gimnasio. Entiendo que este dato es considerado sensible, "
        "que su entrega es voluntaria, y que puedo solicitar su eliminación en cualquier momento "
        "dirigiéndome al gimnasio, sin que ello afecte otros aspectos de mi membresía salvo el uso del "
        "acceso biométrico.")
    add_para(doc, "Autorizo:   [  ] Sí       [  ] No (en caso de \"No\", el acceso se validará por otro medio disponible, por ejemplo código QR).")

    doc.add_paragraph()
    p = add_para(doc, "Firma del socio (mayor de edad): _____________________________", bold=True)

    add_heading(doc, "4. Sección exclusiva para menores de edad")
    add_para(doc,
        "En caso de que el socio sea menor de 18 años, el tratamiento de sus datos personales, "
        "incluyendo el dato biométrico, requiere el consentimiento expreso de su padre, madre o "
        "tutor legal, quien deberá completar y firmar esta sección.")
    add_datos_table(doc, [
        ("Nombre del padre/madre/tutor", "_______________________________________________"),
        ("Parentesco", "_______________________________________________"),
        ("Identificación oficial (tipo y número)", "_______________________________________________"),
        ("Teléfono de contacto del tutor", "_______________________________________________"),
    ])
    add_para(doc,
        "Declaro ser el padre, madre o tutor legal del menor registrado en este formato, y autorizo el "
        "tratamiento de sus datos personales, incluyendo su huella dactilar para control de acceso, "
        "conforme a lo descrito en las secciones 2 y 3 de este documento.")
    add_para(doc, "Autorizo:   [  ] Sí       [  ] No")
    doc.add_paragraph()
    add_para(doc, "Firma del padre/madre/tutor: _____________________________", bold=True)

    add_heading(doc, "5. Derechos del titular")
    add_para(doc,
        f"El titular (o su tutor, en caso de menores) puede ejercer sus derechos de Acceso, "
        f"Rectificación, Cancelación u Oposición (ARCO), así como revocar este consentimiento en "
        f"cualquier momento, dirigiéndose directamente a {CLIENTE} a través de su representante "
        f"{REP_CLIENTE}.")

    add_datos_table(doc, [
        ("Gimnasio", CLIENTE),
        ("Fecha", "___ /___ /______"),
        ("Lugar", ""),
    ])

    add_footer_line(doc)
    doc.save(f"{OUT_DIR}\\GymAccess_Formato_Inscripcion_Consentimiento_Socio.docx")
    print("Formato de Inscripcion y Consentimiento generado.")


if __name__ == "__main__":
    gen_aviso_privacidad()
    gen_nda()
    gen_contrato_servicios()
    gen_consentimiento_socio()
